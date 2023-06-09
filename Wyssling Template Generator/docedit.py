# --FUTURE IDEAS--
# handle templates for multiple states
# handle templates for different types of mounts (ground, ballast, etc)
# link with database to autofill info for states
# make compatible with mac as well? (maybe not necessary if using on windows server)



from docx2pdf import convert
from datetime import date
import os
from docx import Document #make sure this is something that can be used on other PCs

#template = Document('NJ Template.docx')
#template = Document('NJ Template (with defVals).docx')
global template
#template_copy = Document('NJ Template (with defVals).docx')

#fill values

date_fill = "March 7, 2023"
name_fill = "Cheeser"
address_fill = "324 Congress Street, Saddle Brook"
state_fill = "NJ"
systemSize_fill = "9.600" 
framing_fill = "2x6 dimensional lumber at 16” on center."
roofMaterial_fill = "Composite Asphalt Shingles"
slope_fill = "26" 
atticAccess_fill = "Accessible" 
existingDead_fill = "7" 
newDead_fill = "3"
totalDead_fill = "10"
snow_fill = "30"
asce_fill = "16"
wind_fill = "120" 
codeYear_fill = "2021" 
exposureCat_fill = "B" 
mountingType_fill = "EcoFasten" 
mountingInfo_fill = "The maximum allowable withdrawal force for a 5/16” lag screw is 229 lbs per inch of penetration as identified in the National Design Standards (NDS) of timber construction specifications. Based on a minimum penetration depth of 2½”, the allowable capacity per connection is greater than the design withdrawal force (demand). Considering the variable factors for the existing roof framing and installation tolerances, the connection using one 5/16” diameter lag screw with a minimum of 2½” embedment will be adequate and will include a sufficient factor of safety."
spacing_fill = "48"

#default vals override with defvals doc

date_default = "dateDef"
name_default = "nameDef"
address_default = "addressDef"
state_default = "stateDef"
systemSize_default = "systemDef" 
framing_default = "framingDef"
roofMaterial_default = "materialDef"
slope_default = "slopeDef" 
atticAccess_default = "accessDef" 
existingDead_default = "existDef" 
newDead_default = "newDef"
totalDead_default = "totalDef"
snow_default = "snowDef" 
asce_default = "asceDef"
wind_default = "windDef" 
codeYear_default = "yearDef" 
exposureCat_default = "exposDef" 
mountingType_default = "mountDef" 
mountingInfo_default = "mountInfoDef"
spacing_default = "spacingDef"

#userInput values

name = "Tester"
address = "Testington Road, Testbrook"
state = "TE"
systemSize = "6.900" #use float or string?
framing = "6x9 dimensional lumber at 69” on center."
roofMaterial = "Cool Awesome Sauce"
slope = "420" #use int/float or string?
atticAccess = "Testable" #use boolean?
existingDead = "69" #use float or string?
newDead = "666"
totalDead = "1337"
snow = "33" #set default values for state
asce = "88"
wind = "999" #set default values for state
codeYear = "2018" #make int? 
exposureCat = "F" #make char?
mountingType = "NiceRack" 
#changes based on the mounting the user selects
mountingInfo = "The maximum allowable withdrawal force for a 5/16” lag screw is 229 lbs per inch of penetration as identified in the National Design Standards (NDS) of timber construction specifications. Based on a minimum penetration depth of 2½”, the allowable capacity per connection is greater than the design withdrawal force (demand). Considering the variable factors for the existing roof framing and installation tolerances, the connection using one 5/16” diameter lag screw with a minimum of 2½” embedment will be adequate and will include a sufficient factor of safety."
spacing = "51"

def setUserInput(userInput): #TODO: figure out better system than hard coding values
    global name
    name = userInput[0]

    global address
    address = userInput[1]

    global state
    state = userInput[15]

    global systemSize
    systemSize = userInput[2]

    global framing
    framing = userInput[3]

    global roofMaterial
    roofMaterial = userInput[4]

    global slope
    slope = userInput[5]

    global atticAccess
    atticAccess = userInput[14]

    global existingDead
    existingDead = userInput[6]

    global newDead
    newDead = userInput[7]

    global totalDead
    totalDead = str(float(newDead) + float(existingDead))

    global snow
    snow = userInput[8]

    global asce
    asce = userInput[16]

    global wind
    wind = userInput[9]

    global codeYear
    codeYear = userInput[10]

    global exposureCat
    exposureCat = userInput[17]

    global mountingType
    mountingType = userInput[11]

    global mountingInfo
    mountingInfo = userInput[12]

    global spacing
    spacing = userInput[13]

    #print(userInput)




today = date.today()
letterDate = today.strftime("%B %d, %Y")


def replace(default, replace, paragraphIn):
    for run in paragraphIn.runs:
        if default in run.text:
            print(default + " " + replace)
            run.text = run.text.replace(default, replace)


def run():
    #template = Document('NJ Template (with defVals).docx')
    template = Document('templates/NJ Template (with defVals).docx')
    for paragraph in template.paragraphs:
        #print(paragraph.text)
        if not paragraph.text:
            continue

        #date
        if date_default in paragraph.text:
            replace(date_default, letterDate, paragraph)

        #name
        if name_default in paragraph.text:
            replace(name_default, name, paragraph)

        #address
        if address_default in paragraph.text:
            replace(address_default, address, paragraph)
            
        #state
        if state_default in paragraph.text:
            replace(state_default, state, paragraph)

        #systemSize
        if systemSize_default in paragraph.text:
            replace(systemSize_default, systemSize, paragraph)
        
        #framing
        if framing_default in paragraph.text:
            replace(framing_default, framing, paragraph)
            

        #roofMaterial
        if roofMaterial_default in paragraph.text:
            replace(roofMaterial_default, roofMaterial, paragraph)

        #slope
        if slope_default in paragraph.text:
            replace(slope_default, slope, paragraph)

        #atticAccess
        if atticAccess_default in paragraph.text:
            replace(atticAccess_default, atticAccess, paragraph)

        #existingDead
        if existingDead_default in paragraph.text:
            replace(existingDead_default, existingDead, paragraph)

        #newDead
        if newDead_default in paragraph.text:
            replace(newDead_default, newDead, paragraph)

        #totalDead
        if totalDead_default in paragraph.text:
            replace(totalDead_default, totalDead, paragraph)

        #snow
        if snow_default in paragraph.text:
            replace(snow_default, snow, paragraph)

        #asce
        if asce_default in paragraph.text:
            replace(asce_default, asce, paragraph)

        #wind
        if wind_default in paragraph.text:
            replace(wind_default, wind, paragraph)

        #exposure category
        if exposureCat_default in paragraph.text:
            replace(exposureCat_default, exposureCat, paragraph)

        #code year
        if codeYear_default in paragraph.text:
            replace(codeYear_default, codeYear, paragraph)

        #mount type
        if mountingType_default in paragraph.text:
            replace(mountingType_default, mountingType, paragraph)

        #mount info
        if mountingInfo_default in paragraph.text:
            replace(mountingInfo_default, mountingInfo, paragraph)

        #spacing
        if spacing_default in paragraph.text:
            replace(spacing_default, spacing, paragraph)
        

    fileDate = today.strftime("%m-%d-%Y")
    filename = name + " " + fileDate
    #template.save(filename + '.docx')
    template.save('generated/' + filename + '.docx')
    
    #convert(filename + '.docx', filename + ' -C' + '.pdf')
    convert('generated/' + filename + '.docx', 'generated/' + filename + ' -C' + '.pdf')

    try:
        os.remove('generated/' + filename + '.docx')
        print("File deleted successfully")
    except OSError as e:
        print(f"Error deleting file: {e}")

#template = Document('NJ Template (with defVals).docx')
template = Document('templates/NJ Template (with defVals).docx')

#run()

